#-*- coding:utf-8 -*-
import os
import docx
from time import sleep
from tqdm import tqdm
from pyltp import SentenceSplitter
from pyltp import Segmentor
from pyltp import Postagger
import tkinter as tk
from tkinter import filedialog
import docx
import docx2python
# from win32com import client as wc
import xlsxwriter as xw

delthing = '/nh'+"犯"+"/n"+"/v"+"/ns"+'/p'+'/d'+'/i'+'/r'+'/a'+'/nz'+'/c'+'/m'+'/nt'+'/b'
root = tk.Tk()
root.withdraw()
print("选择文档所在的文件夹后，请耐心等待。文档数量越大、doc数量越多，处理速度将会延长。")
print("初始数据一定要做好备份！")
fn = filedialog.askdirectory()      # 打开文件夹选择框


morethanone = "被害人数量超过1人,目前尚无法自动提取，请人工核查。"
tiqu_fail = []
tiqu_fail_quchong = []
zhuanhuan_fail = []
bgrname_str = ""

path = fn+'\\'
resultlist = []

def xw_toExcel(data, fileName):  # xlsxwriter库储存数据到excel
    workbook = xw.Workbook(fileName)  # 创建工作簿
    worksheet1 = workbook.add_worksheet("sheet1")  # 创建子表
    worksheet1.activate()  # 激活表
    title = ['案卷名称', '被告人', '被害人',"前科"]  # 设置表头
    worksheet1.write_row('A1', title)  # 从A1单元格开始写入表头
    i = 2  # 从第二行开始写入数据
    for j in range(len(data)):
        insertData = [data[j]["案卷名称"], data[j]["被告人"], data[j]["被害人"], data[j]["前科"]]
        row = 'A' + str(i)
        worksheet1.write_row(row, insertData)
        i += 1
    workbook.close()  # 关闭表

'''
def save_doc_to_docx(rawpath):  # doc转docx

    # :param rawpath: 传入和传出文件夹的路径
    # :return: None

    word = wc.Dispatch("Word.Application")
    # 不能用相对路径，老老实实用绝对路径
    # 需要处理的文件所在文件夹目录
    filenamelist = os.listdir(rawpath)
    for i in filenamelist:
        # 找出文件中以.doc结尾并且不以~$开头的文件（~$是为了排除临时文件的）
        if i.endswith('.doc') and not i.startswith('~$'):
            try: # try
                # 打开文件
                doc = word.Documents.Open(rawpath + i)
                # # 将文件名与后缀分割
                rename = os.path.splitext(i)
                # 将文件另存为.docx
                doc.SaveAs(path + rename[0] + '.docx', 12)  # 12表示docx格式
                doc.Close
            except:
                zhuanhuan_fail.append(i)
    word.Quit()
'''
def save_doc_to_docx(folder_path):
    for filename in os.listdir(folder_path):
        if filename.endswith('.doc'):
            doc_file = os.path.join(folder_path, filename)
            docx_file = os.path.join(folder_path, f'{os.path.splitext(filename)[0]}.docx')
            doc = docx2python.docx2python(doc_file)
            docx_file = docx.Document()
            for para in doc.body:
                docx_file.add_paragraph(para)
            docx_file.save(docx_file)
save_doc_to_docx(path)

'''获取文件夹下所有docx文档'''
def get_all_file(dir_name):
    filename_list = []
    for root, dirs, files in os.walk(dir_name):
        for filename in files:
            # 文件名列表，包含完整路径
            #fullname_list.append(os.path.join(root, filename))
            # 文件名列表，只包含文件名
            if filename.endswith('.docx'):         #判断是否为docx文档
                filename_list.append(filename)
    return filename_list

all_files = get_all_file(fn)    #全部的文档列表
len_all_files = len(all_files)      #获取文档列表长度

#fn = r'E:\codeFIle\files\中华人民共和国最高人民法院2.docx'      #文档绝对路径
'''加载pyltp的模型'''
lexicon_path = '/Users/xumuzhi/coding/nameGet/lexicon.txt'
LTP_DATA_DIR = r'/Users/xumuzhi/coding/nameGet/ltp_data_v3.4.0'      # ltp模型目录的路径
ner_model_path = os.path.join(LTP_DATA_DIR, 'ner.model')      # 命名实体识别模型路径，模型名称为`ner.model`
cws_model_path = os.path.join(LTP_DATA_DIR, 'cws.model')      # 命名实体识别模型路径，模型名称为`cws.model`
pos_model_path = os.path.join(LTP_DATA_DIR, 'pos.model')      # 命名实体识别模型路径，模型名称为`pos.model`


k = 0      #赋初值

'''分词'''
def segmentor(sentence=''):
    segmentor=Segmentor()#初始化实例
    segmentor.load_with_lexicon(cws_model_path,lexicon_path)  # 加载模型
    words=segmentor.segment(sentence)#产生分词
    words_list=list(words)
    segmentor.release()#释放模型
    return words_list

'''词性标注'''
def posttagger(words):
    postagger=Postagger()#初始化实例
    postagger.load_with_lexicon(pos_model_path,lexicon_path)
    postags=postagger.postag(words)#词性标注
    postagger.release()
    return postags

'''获取被告人姓名'''
def BGRnameGet(bgr,files):
    global bgrname
    bgrname = ""
    i = 0      #赋初值
    j = 0      #赋初值
    k = 0      #赋初值
    m = 0      #赋初值
    list1 = []      #创建列表1用于存放文章分段  
    list2 = []      #创建列表2用于存放包含被告人的分段
    list3 = []      #创建列表3用于存放词性为人名的词语
    # bgrname_list = []       #创建被告人列表用于去重
    bgrname_quchong = []       #创建去重被告人列表
    
    try:
        doc = docx.Document(fn+'/'+files)      #读取文档       
        for paragraph in doc.paragraphs:      #把文章分段加入列表
            list1.append(paragraph.text)   

        len_list1 = len(list1)      #统计列表长度
        for i in range(0,int(len_list1)):      #将开头为报告人的段落输入列表2
            if list1[i].startswith(bgr):
                list2.append(list1[i])
            
        len_list2 = len(list2)
        for j in  range(0,len_list2):
            words=segmentor(sentence= list2[j])
            postags=posttagger(words)   
            for word,tag in zip(words,postags):
                #print(word+'/'+tag)
                list3.append(word+'/'+tag)
        '''测试'''
        # print(list3)

        len_list3 = len(list3)  
        '''      
        for k in range(0,len_list3):              
            if (list3[k].find(bgr) != -1) and (list3[k+1].find("/nh") != -1) and (list3[k+2].find("wp") != -1):         
                bgrname = bgrname.strip('犯')+"、"+list3[k+1].strip('/nh'+"犯"+"/n"+"/v"+"/ns"+'/p'+'/d'+'/i'+'/r'+'/a'+'/nz')
            elif (list3[k].find(bgr) != -1) and (list3[k+1].find("/nh") != -1) and (list3[k+3].find("wp") != -1) and not ((list3[k+2].find("罪") != -1)):
                bgrname = bgrname.strip('犯')+"、"+list3[k+1].strip('/nh'+"犯"+"/n"+"/v"+"/ns"+'/p'+'/d'+'/i'+'r'+'/a'+'/nz')+list3[k+2].strip('/nh'+"犯"+"/n"+"/v"+"/ns"+'/p'+'/d'+'/i'+'r'+'/a'+'/nz')
        '''
        for k in range(0,len_list3):     # 后边儿是性别的         
            if (list3[k].find(bgr) != -1) and (list3[k+2].find("，") != -1) and (list3[k+3].find("/b") != -1) and not (list3[k+1].find("罪") != -1):         
                bgrname = bgrname.strip('犯')+"、"+list3[k+1].strip(delthing)
            elif (list3[k].find(bgr) != -1) and (list3[k+3].find("，") != -1) and (list3[k+4].find("/b") != -1) and not (list3[k+2].find("罪") != -1):         
                bgrname = bgrname.strip('犯')+"、"+list3[k+1].strip(delthing)+list3[k+2].strip(delthing)
            elif (list3[k].find(bgr) != -1) and (list3[k+4].find("，") != -1) and (list3[k+5].find("/b") != -1) and not (list3[k+3].find("罪") != -1):         
                bgrname = bgrname.strip('犯')+"、"+list3[k+1].strip(delthing)+list3[k+2].strip(delthing)+list3[k+3].strip(delthing) 

        for k in range(0,len_list3):    # 后边儿是“曾”的            
            if (list3[k].find(bgr) != -1) and (list3[k+2].find("，") != -1) and (list3[k+3].find("曾/d") != -1) and not (list3[k+1].find("罪") != -1):         
                bgrname = bgrname.strip('犯')+"、"+list3[k+1].strip(delthing)
            elif (list3[k].find(bgr) != -1) and (list3[k+3].find("，") != -1) and (list3[k+4].find("曾/d") != -1) and not (list3[k+2].find("罪") != -1):         
                bgrname = bgrname.strip('犯')+"、"+list3[k+1].strip(delthing)+list3[k+2].strip(delthing)
            elif (list3[k].find(bgr) != -1) and (list3[k+4].find("，") != -1) and (list3[k+5].find("曾/d") != -1) and not (list3[k+3].find("罪") != -1):         
                bgrname = bgrname.strip('犯')+"、"+list3[k+1].strip(delthing)+list3[k+2].strip(delthing)+list3[k+3].strip(delthing)
        
        for k in range(0,len_list3):    # 后边儿是“（曾”的  and (list3[k+3].find("曾") != -1)          
            if (list3[k].find(bgr) != -1) and (list3[k+2].find("（") != -1)  and not (list3[k+1].find("罪") != -1):         
                bgrname = bgrname.strip('犯')+"、"+list3[k+1].strip(delthing)
            elif (list3[k].find(bgr) != -1) and (list3[k+3].find("（") != -1) and not (list3[k+2].find("罪") != -1):         
                bgrname = bgrname.strip('犯')+"、"+list3[k+1].strip(delthing)+list3[k+2].strip(delthing)
            elif (list3[k].find(bgr) != -1) and (list3[k+4].find("（") != -1) and not (list3[k+3].find("罪") != -1):         
                bgrname = bgrname.strip('犯')+"、"+list3[k+1].strip(delthing)+list3[k+2].strip(delthing)+list3[k+3].strip(delthing)

        for k in range(0,len_list3):    # 后边儿是“绰号”的            
            if (list3[k].find(bgr) != -1) and (list3[k+2].find("，") != -1) and (list3[k+3].find("绰") != -1) and not (list3[k+1].find("罪") != -1):         
                bgrname = bgrname.strip('犯')+"、"+list3[k+1].strip(delthing)
            elif (list3[k].find(bgr) != -1) and (list3[k+3].find("，") != -1) and (list3[k+4].find("绰") != -1) and not (list3[k+2].find("罪") != -1):         
                bgrname = bgrname.strip('犯')+"、"+list3[k+1].strip(delthing)+list3[k+2].strip(delthing)
            elif (list3[k].find(bgr) != -1) and (list3[k+4].find("，") != -1) and (list3[k+5].find("绰") != -1) and not (list3[k+3].find("罪") != -1):         
                bgrname = bgrname.strip('犯')+"、"+list3[k+1].strip(delthing)+list3[k+2].strip(delthing)+list3[k+3].strip(delthing)

            # elif or(list3[k+3].find("曾") != -1)}
    except:
        zhuanhuan_fail.append(files)

    bgrname = bgrname.lstrip(" 、")         # 去掉开头的顿号
    bgrname = bgrname.rstrip("犯"+"签署"+"辩解"+"辩称"+"老年体弱"+"到案后"+"归案后"+"表示谅解")         # 再次去掉犯字
    bgrname_list = bgrname.split("、")      # 以顿号分割
    i = 0
    for i in bgrname_list:                  # 列表去重
        if i not in bgrname_quchong:
            bgrname_quchong.append(i)
    bgrname_str = '、'.join(bgrname_quchong)
    return bgrname_str

    

'''获取被害人姓名'''
#-*- coding:utf-8 -*-
def BHRnameGet(bhr,files):
    #-*- coding:utf-8 -*-
    global bhrname
    bhrname = ""
    i = 0      #赋初值
    j = 0      #赋初值
    k = 0      #赋初值
    m = 0      #赋初值
    list1 = []      #创建列表1用于存放文章分段  
    list2 = []      #创建列表2用于存放包含被告人的分段
    list3 = []      #创建列表3用于存放词性为人名的词语
    # bgrname_list = []       #创建被告人列表用于去重
    bhrname_quchong = []       #创建去重被告人列表
    
    try:
        doc = docx.Document(fn+'/'+files)      #读取文档       
        for paragraph in doc.paragraphs:      #把文章分段加入列表
            list1.append(paragraph.text)   

        len_list1 = len(list1)      #统计列表长度
        for i in range(0,int(len_list1)):      #将包含被害人的段落输入列表2
            if list1[i].find(bhr) != -1:
                list2.append(list1[i])
            
        len_list2 = len(list2)
        for j in  range(0,len_list2):
            words=segmentor(sentence= list2[j])
            postags=posttagger(words)   
            for word,tag in zip(words,postags):
                #print(word+'/'+tag)
                list3.append(word+'/'+tag)
        '''测试'''
        # print(list3)
 #  
        len_list3 = len(list3)  
        for k in range(0,len_list3):     #        
            if (list3[k].find(bhr) != -1) and (list3[k+1].endswith("/nh")):  
                if len(list3[k+1]) != 4:     # 如果名字不为一个字
                    bhrname = bhrname.strip('犯')+"、"+list3[k+1].strip(delthing)
                    bhrname = bhrname.rstrip("犯"+"签署"+"辩解"+"辩称"+"老年体弱"+"到案后"+"归案后"+"表示谅解"+'的/u')     # 去掉犯和其他不相关的字
                elif len(list3[k+1]) == 4:      # 如果名字只有一个字
                    bhrname = bhrname.strip(delthing)+"、"+list3[k+1].strip(delthing)+list3[k+2].strip(delthing)
                    bhrname = bhrname.rstrip("犯"+"签署"+"辩解"+"辩称"+"老年体弱"+"到案后"+"归案后"+"表示谅解"+'的/u')     # 去掉犯和其他不相关的字
            
            # elif or(list3[k+3].find("曾") != -1)}
    except:
        tiqu_fail.append(files)

    bhrname = bhrname.lstrip(" 、")         # 去掉开头的顿号
    bhrname = bhrname.rstrip("犯"+"签署"+"辩解"+"辩称"+"老年体弱"+"到案后"+"归案后"+"表示谅解"+'的/u')         # 再次去掉犯和其他不相关的字
    if bhrname == "":
        bhrname = "(无被害人）、(无被害人）"
    bhrname_list = bhrname.split("、")      # 以顿号分割
    i = 0
    for i in bhrname_list:                  # 列表去重
        if i not in bhrname_quchong:
            bhrname_quchong.append(i)
    bhrname_str = '、'.join(bhrname_quchong)
    return bhrname_str

'''获取前科劣迹'''
def QiankeGet(files):
    #-*- coding:utf-8 -*-
    # try:
    print(files)
    i = 0
    list1 = []
    global QKjieguo
    global jiheqilai
    jiheqilai = ""
    try:
        doc = docx.Document(fn+'\\'+files)      #读取文档       
        for paragraph in doc.paragraphs:      #把文章分段加入列表
            list1.append(paragraph.text)

        len_list1 = len(list1) 
        
        for i in range(0,len_list1): 
            if list1[i].startswith('被告人'):
                jiaru = list1[i]
                sents = SentenceSplitter.split(jiaru)  # 分句
                test = ('\n'.join(sents))
                jiheqilai = jiheqilai + test
        array = jiheqilai.split('\n')        # 句子分割加入数组
        # print(array)

        # len_array = len(array)  
        j = 0
        if array[j].startswith('被告人'):
            if array[j].find("因本案")!= -1:
                QKjieguo = "(无前科）"
            elif array[j+1].startswith("因本案"):
                QKjieguo = "(无前科）"
            elif array[j+2].startswith("因本案"):
                QKjieguo = array[j+1]
            elif array[j+3].startswith("因本案"):
                QKjieguo = array[j+1]+array[j+2]
            elif array[j+4].startswith("因本案"):
                QKjieguo = array[j+1]+array[j+2]+array[j+3]
            elif array[j+5].startswith("因本案"):
                QKjieguo = array[j+1]+array[j+2]+array[j+3]+array[j+4]
            elif array[j+6].startswith("因本案"):
                QKjieguo = array[j+1]+array[j+2]+array[j+3]+array[j+4]+array[j+5]
            else:
                QKjieguo = '（无前科）'
        return QKjieguo
    except:
        print(files+'前科提取未知错误')


'''
for k in range(0,len_all_files):
    print("\n")
    print(all_files[k]+":")
    print("被告人："+nameGet("被告人",all_files[k]))
    print("被害人："+nameGet("被害人",all_files[k]))
'''
print("-----------------正在运算中，请稍后-----------------")
print("根据文件数量，运算的时间可能会有较大差别，请耐心等待。")
pbar = tqdm(total = len_all_files)
for k in range(0,len_all_files):
    wcjd = 1/len_all_files
    sysj = (len_all_files-k)*4/60
    '''进度条'''
    sleep(0.01)
    pbar.update(1)
    # print("完成进度:"+str('%.2f'% wcjd)+"%"+";  预计剩余时间："+str('%.1f'% sysj)+"分钟")
    info_Qianke = QiankeGet(all_files[k])
    nameget_bgr = BGRnameGet("被告人",all_files[k])
    nameget_bhr = BHRnameGet("被害人",all_files[k])
    putittoresult = {"案卷名称":all_files[k].strip('.docx'),"被告人":nameget_bgr,"被害人":nameget_bhr,"前科":info_Qianke}
    resultlist.append (putittoresult)
    # print(resultlist[k])
    # tiqu_fail.append(all_files[k])
pbar.close()
print("--------请稍等，正在准备写入表格--------")  
'''输出提取与转化失败的'''
i = 0
for i in tiqu_fail:                  # 列表去重
    if i not in tiqu_fail_quchong:
        tiqu_fail_quchong.append(i)   
print("--------以下文件提取失败，正在准备二次提取--------")   
print(tiqu_fail_quchong)
fileName = r'C:\Users\muzzo\Desktop\提取结果.xlsx'
xw_toExcel(resultlist, fileName)
input("运算完成，请在桌面查看结果。按回车键退出程序")